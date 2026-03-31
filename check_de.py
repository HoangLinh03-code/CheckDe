# -*- coding: utf-8 -*-
"""
Check Đề - Exam Answer Verification Tool
=========================================
So sánh đáp án của các đề trộn với đề gốc, phát hiện sai sót.
Hỗ trợ: docx + pdf, equation/công thức toán học.

Quy trình:
1. Parse đề gốc (docx) → danh sách câu hỏi + đáp án (bao gồm equations)
2. Parse các đề trộn (docx/pdf) → danh sách câu hỏi + đáp án
3. Đọc bảng đáp án Excel
4. Dùng AI (Gemini) + gửi file gốc để khớp câu hỏi (hỗ trợ equations)
5. Xác minh đáp án
6. Kiểm tra cấu trúc/định dạng
7. Xuất báo cáo Excel
"""

import os
import sys
import re
import json
import argparse
import time
import glob
from datetime import datetime
import traceback
from lxml import etree

# Đảm bảo encoding UTF-8
if hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

# PDF support - try multiple libraries
try:
    import fitz  # PyMuPDF
    PDF_LIB = 'fitz'
    print("[Init] PDF support: PyMuPDF (fitz)")
except ImportError:
    try:
        import pdfplumber
        PDF_LIB = 'pdfplumber'
        print("[Init] PDF support: pdfplumber")
    except ImportError:
        PDF_LIB = None
        print("[Init] ⚠️  Không có thư viện PDF (pip install PyMuPDF hoặc pdfplumber)")
        print("       File PDF sẽ được gửi trực tiếp đến AI để phân tích.")

# Import callAPI từ cùng thư mục
sys.path.append(os.path.dirname(os.path.abspath(__file__)))
from callAPI import get_vertex_ai_credentials, VertexClient

# Import math_exam_handler
try:
    from math_exam_handler import (
        detect_subject,
        process_math_exam_v2,
        normalize_math_answer,
        answers_match as math_answers_match,
        verify_math_answers,
        docx_has_wmf_equations,
        MathQuestion,
    )
    MATH_HANDLER_AVAILABLE = True
except ImportError:
    MATH_HANDLER_AVAILABLE = False
    print("[Init] ⚠️  math_exam_handler không khả dụng")

# Import prompt_loader
try:
    from prompt_loader import (
        build_prompt_header as load_prompt_header,
        get_output_format as load_output_format,
        get_vision_prompt as load_vision_prompt,
        get_max_pages as load_max_pages,
        get_subject_label,
    )
    PROMPT_LOADER_AVAILABLE = True
except ImportError:
    PROMPT_LOADER_AVAILABLE = False
    print("[Init] ⚠️  prompt_loader không khả dụng")

# OMML namespace cho trích xuất equation
OMML_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
WML_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


# ============================================================
# 1. DOCX PARSER - Bóc tách câu hỏi từ file Word
# ============================================================

class Question:
    """Đại diện cho một câu hỏi trắc nghiệm."""
    def __init__(self, number, text, options):
        self.number = number        # Số thứ tự câu (int)
        self.text = text            # Nội dung câu hỏi (str)
        self.options = options      # Dict {'A': 'nội dung A', 'B': '...', ...}

    def __repr__(self):
        opts = ', '.join(f"{k}: {v[:30]}..." for k, v in self.options.items())
        return f"Q{self.number}: {self.text[:50]}... | {opts}"


def normalize_text(text):
    """Chuẩn hóa text để so sánh: bỏ khoảng trắng thừa, lowercase."""
    text = re.sub(r'\s+', ' ', text).strip()
    return text


def omml_to_text(omath_element):
    """
    Chuyển đổi OMML (Office Math Markup Language) element thành text đọc được.
    Xử lý: phân số, lũy thừa, căn, chỉ số, ký hiệu đặc biệt.
    """
    result = []
    
    for child in omath_element:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        
        if tag == 'r':  # Run - text thường hoặc ký hiệu
            for t in child.iter():
                t_tag = t.tag.split('}')[-1] if '}' in t.tag else t.tag
                if t_tag == 't' and t.text:
                    result.append(t.text)
        
        elif tag == 'f':  # Fraction (phân số)
            num_el = child.find(f'{{{OMML_NS}}}num')
            den_el = child.find(f'{{{OMML_NS}}}den')
            num_text = omml_to_text(num_el) if num_el is not None else '?'
            den_text = omml_to_text(den_el) if den_el is not None else '?'
            result.append(f'({num_text}/{den_text})')
        
        elif tag == 'rad':  # Radical (căn)
            deg_el = child.find(f'{{{OMML_NS}}}deg')
            e_el = child.find(f'{{{OMML_NS}}}e')
            e_text = omml_to_text(e_el) if e_el is not None else '?'
            if deg_el is not None and omml_to_text(deg_el).strip():
                deg_text = omml_to_text(deg_el)
                result.append(f'√[{deg_text}]({e_text})')
            else:
                result.append(f'√({e_text})')
        
        elif tag == 'sSup':  # Superscript (lũy thừa)
            base_el = child.find(f'{{{OMML_NS}}}e')
            sup_el = child.find(f'{{{OMML_NS}}}sup')
            base_text = omml_to_text(base_el) if base_el is not None else '?'
            sup_text = omml_to_text(sup_el) if sup_el is not None else '?'
            result.append(f'{base_text}^{sup_text}')
        
        elif tag == 'sSub':  # Subscript (chỉ số dưới)
            base_el = child.find(f'{{{OMML_NS}}}e')
            sub_el = child.find(f'{{{OMML_NS}}}sub')
            base_text = omml_to_text(base_el) if base_el is not None else '?'
            sub_text = omml_to_text(sub_el) if sub_el is not None else '?'
            result.append(f'{base_text}_{sub_text}')
        
        elif tag == 'sSubSup':  # Sub+Superscript
            base_el = child.find(f'{{{OMML_NS}}}e')
            sub_el = child.find(f'{{{OMML_NS}}}sub')
            sup_el = child.find(f'{{{OMML_NS}}}sup')
            base_text = omml_to_text(base_el) if base_el is not None else '?'
            sub_text = omml_to_text(sub_el) if sub_el is not None else '?'
            sup_text = omml_to_text(sup_el) if sup_el is not None else '?'
            result.append(f'{base_text}_{sub_text}^{sup_text}')
        
        elif tag == 'd':  # Delimiter (ngoặc)
            e_elements = child.findall(f'{{{OMML_NS}}}e')
            inner = '; '.join(omml_to_text(e) for e in e_elements)
            result.append(f'({inner})')
        
        elif tag == 'nary':  # N-ary (tổng, tích phân)
            sub_el = child.find(f'{{{OMML_NS}}}sub')
            sup_el = child.find(f'{{{OMML_NS}}}sup')
            e_el = child.find(f'{{{OMML_NS}}}e')
            # Lấy ký hiệu
            chr_el = child.find(f'{{{OMML_NS}}}naryPr/{{{OMML_NS}}}chr')
            sym = '∑'
            if chr_el is not None:
                sym = chr_el.get(f'{{{OMML_NS}}}val', '∑')
            e_text = omml_to_text(e_el) if e_el is not None else '?'
            result.append(f'{sym}({e_text})')
        
        elif tag == 'func':  # Function (sin, cos, log...)
            fname_el = child.find(f'{{{OMML_NS}}}fName')
            e_el = child.find(f'{{{OMML_NS}}}e')
            fname_text = omml_to_text(fname_el) if fname_el is not None else 'f'
            e_text = omml_to_text(e_el) if e_el is not None else '?'
            result.append(f'{fname_text}({e_text})')
        
        elif tag == 'limLow' or tag == 'limUpp':  # Limit
            e_el = child.find(f'{{{OMML_NS}}}e')
            lim_el = child.find(f'{{{OMML_NS}}}lim')
            e_text = omml_to_text(e_el) if e_el is not None else '?'
            lim_text = omml_to_text(lim_el) if lim_el is not None else '?'
            result.append(f'{e_text}_({lim_text})')
        
        elif tag == 'e' or tag == 'num' or tag == 'den' or tag == 'deg' or tag == 'sub' or tag == 'sup' or tag == 'fName' or tag == 'lim':
            # Container elements - recurse
            result.append(omml_to_text(child))
        
        elif tag == 'oMath':  # Nested oMath
            result.append(omml_to_text(child))
        
        else:
            # Fallback: try to get any text content
            inner_text = omml_to_text(child)
            if inner_text:
                result.append(inner_text)
    
    return ''.join(result)


def extract_paragraph_text_with_equations(para_element):
    """
    Trích xuất text từ một paragraph element, bao gồm cả equations (OMML).
    Trả về chuỗi text với equations được chuyển đổi thành dạng đọc được.
    """
    parts = []
    
    for child in para_element:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        
        if tag == 'r':  # Run - text thường
            for t in child.iter():
                t_tag = t.tag.split('}')[-1] if '}' in t.tag else t.tag
                if t_tag == 't' and t.text:
                    parts.append(t.text)
        
        elif tag == 'oMathPara' or tag == 'oMath':  # Equation
            eq_text = omml_to_text(child)
            if eq_text.strip():
                parts.append(f' {eq_text} ')
        
        elif tag == 'hyperlink':  # Hyperlink
            for t in child.iter():
                t_tag = t.tag.split('}')[-1] if '}' in t.tag else t.tag
                if t_tag == 't' and t.text:
                    parts.append(t.text)
    
    return ''.join(parts).strip()


def extract_all_text_from_docx(filepath):
    """
    Trích xuất TOÀN BỘ text từ file docx, bao gồm equations và tables.
    Trả về danh sách các dòng text.
    """
    doc = Document(filepath)
    lines = []
    
    for element in doc.element.body:
        tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag
        
        if tag == 'p':  # Paragraph
            text = extract_paragraph_text_with_equations(element)
            if text.strip():
                lines.append(text.strip())
        
        elif tag == 'tbl':  # Table
            # Lấy text từ mỗi cell (bao gồm equations)
            for tc in element.iter():
                tc_tag = tc.tag.split('}')[-1] if '}' in tc.tag else tc.tag
                if tc_tag == 'tc':
                    cell_parts = []
                    for p in tc:
                        p_tag = p.tag.split('}')[-1] if '}' in p.tag else p.tag
                        if p_tag == 'p':
                            p_text = extract_paragraph_text_with_equations(p)
                            if p_text.strip():
                                cell_parts.append(p_text.strip())
                    if cell_parts:
                        lines.append(' '.join(cell_parts))
    
    # Fallback
    if not lines:
        for para in doc.paragraphs:
            if para.text.strip():
                lines.append(para.text.strip())
    
    return lines


def extract_text_from_pdf(filepath):
    """
    Trích xuất text từ file PDF.
    Hỗ trợ PyMuPDF (fitz) hoặc pdfplumber.
    """
    lines = []
    
    if PDF_LIB == 'fitz':
        try:
            pdf = fitz.open(filepath)
            for page in pdf:
                text = page.get_text()
                for line in text.split('\n'):
                    if line.strip():
                        lines.append(line.strip())
            pdf.close()
        except Exception as e:
            print(f"  ⚠️  Lỗi đọc PDF bằng PyMuPDF: {e}")
    
    elif PDF_LIB == 'pdfplumber':
        try:
            with pdfplumber.open(filepath) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        for line in text.split('\n'):
                            if line.strip():
                                lines.append(line.strip())
        except Exception as e:
            print(f"  ⚠️  Lỗi đọc PDF bằng pdfplumber: {e}")
    
    if not lines:
        print(f"  ⚠️  Không thể trích xuất text từ PDF: {os.path.basename(filepath)}")
        print(f"       File sẽ được gửi trực tiếp đến AI để phân tích.")
    
    return lines


def get_pdf_page_count(filepath):
    """Đếm số trang PDF."""
    if PDF_LIB == 'fitz':
        try:
            pdf = fitz.open(filepath)
            count = len(pdf)
            pdf.close()
            return count
        except:
            pass
    elif PDF_LIB == 'pdfplumber':
        try:
            with pdfplumber.open(filepath) as pdf:
                return len(pdf.pages)
        except:
            pass
    return 0


def parse_questions_from_text(full_text, source_name=''):
    """
    Parse text để trích xuất danh sách câu hỏi và đáp án.
    Dùng chung cho cả docx và pdf.
    
    Hỗ trợ format:
    - "Câu X:" / "Câu X." (đề tiếng Việt)
    - "Question X:" / "Question X." (đề tiếng Anh)  
    - Đáp án: "A.", "B.", "C.", "D."
    - Options có thể trên cùng dòng hoặc dòng riêng
    
    Returns:
        list[Question]: Danh sách câu hỏi đã parse
    """
    
    # Pattern nhận diện câu hỏi
    # Hỗ trợ: "Câu 1:", "Câu 1.", "Question 1:", "Question 1."
    question_pattern = re.compile(
        r'(?:Câu|Question)\s*(\d+)\s*[.:]\s*',
        re.IGNORECASE
    )
    
    # Pattern nhận diện đáp án trên dòng riêng: "A. xxx"
    # Cũng match khi có tab/spaces phía trước
    option_standalone_pattern = re.compile(
        r'^\s*([A-D])\.\s+(.+)',
        re.IGNORECASE | re.MULTILINE
    )
    
    # Pattern nhận diện đáp án trên cùng dòng (tab-separated): "A. xxx\tB. yyy\tC. zzz\tD. www"
    # Hoặc space-separated: "A. xxx     B. yyy     C. zzz     D. www"
    option_inline_pattern = re.compile(
        r'([A-D])\.\s+(.+?)(?=\s+[A-D]\.\s|$)',
        re.IGNORECASE
    )
    
    questions = []
    
    # Tìm tất cả vị trí câu hỏi
    question_matches = list(question_pattern.finditer(full_text))
    
    if not question_matches:
        print(f"  ⚠️  Không tìm thấy câu hỏi nào trong {source_name}")
        return questions
    
    for i, match in enumerate(question_matches):
        q_num = int(match.group(1))
        q_start = match.end()
        
        # Phạm vi text của câu hỏi này (đến câu tiếp theo hoặc hết văn bản)
        if i + 1 < len(question_matches):
            q_end = question_matches[i + 1].start()
        else:
            q_end = len(full_text)
        
        q_block = full_text[q_start:q_end].strip()
        
        # Tách phần câu hỏi và phần đáp án
        options = {}
        
        # Thử tìm options standalone trước (mỗi option một dòng riêng)
        opt_matches = list(option_standalone_pattern.finditer(q_block))
        
        if opt_matches:
            # Phần câu hỏi = text trước đáp án đầu tiên
            q_text = q_block[:opt_matches[0].start()].strip()
            
            for j, opt_match in enumerate(opt_matches):
                letter = opt_match.group(1).upper()
                
                # Lấy toàn bộ text từ "A. xxx" đến trước "B. yyy"
                if j + 1 < len(opt_matches):
                    full_opt = q_block[opt_match.start():opt_matches[j+1].start()].strip()
                else:
                    full_opt = q_block[opt_match.start():].strip()
                
                # Bỏ prefix "A. "
                full_opt = re.sub(r'^[A-D]\.\s*', '', full_opt).strip()
                options[letter] = full_opt
        else:
            # Thử tìm options inline (tất cả trên cùng dòng)
            # Tìm dòng chứa nhiều option
            for line in q_block.split('\n'):
                inline_matches = list(option_inline_pattern.finditer(line))
                if len(inline_matches) >= 3:  # Ít nhất 3 options trên cùng dòng
                    for m in inline_matches:
                        letter = m.group(1).upper()
                        options[letter] = m.group(2).strip()
                    # Phần câu hỏi = text trước dòng chứa options
                    q_text = q_block[:q_block.index(line)].strip()
                    break
            else:
                q_text = q_block
        
        # Nếu vẫn không tìm thấy đủ options → thử parse linh hoạt hơn
        if len(options) < 4:
            # Reset và thử pattern rộng hơn trên toàn bộ q_block
            alt_options = {}
            alt_pattern = re.compile(r'([A-D])\.\s*(.+?)(?=\s*[A-D]\.\s|$)', re.DOTALL)
            alt_matches = list(alt_pattern.finditer(q_block))
            
            # Chỉ lấy 4 matches cuối cùng (options thường ở cuối block)
            if len(alt_matches) >= 4:
                for m in alt_matches[-4:]:
                    letter = m.group(1).upper()
                    alt_options[letter] = m.group(2).strip().split('\n')[0].strip()
                
                if len(alt_options) >= len(options):
                    options = alt_options
                    # Phần câu hỏi = text trước option đầu tiên trong 4 options cuối
                    first_opt_pos = alt_matches[-4].start()
                    q_text = q_block[:first_opt_pos].strip()
        
        questions.append(Question(q_num, q_text, options))
    
    return questions


def parse_docx_questions(filepath):
    """Parse file docx (có hỗ trợ equations/OMML)."""
    lines = extract_all_text_from_docx(filepath)
    full_text = '\n'.join(lines)
    return parse_questions_from_text(full_text, os.path.basename(filepath))


def parse_pdf_questions(filepath):
    """Parse file PDF."""
    lines = extract_text_from_pdf(filepath)
    if not lines:
        return []  # Sẽ fallback sang AI-based parsing
    full_text = '\n'.join(lines)
    return parse_questions_from_text(full_text, os.path.basename(filepath))


def parse_exam_file(filepath):
    """Parse file đề thi (tự động nhận diện docx/pdf)."""
    ext = os.path.splitext(filepath)[1].lower()
    if ext == '.pdf':
        return parse_pdf_questions(filepath)
    else:  # .docx
        return parse_docx_questions(filepath)


# ============================================================
# 2. EXCEL ANSWER KEY PARSER
# ============================================================

def parse_answer_key_excel(filepath):
    """
    Parse file Excel đáp án.
    
    Cấu trúc Excel:
    Câu | Đề gốc | Mã 0101 | Mã 0102 | ...
     1  |   B    |   A     |   D     | ...
     2  |   B    |   B     |   A     | ...
    
    Returns:
        dict: {
            'gốc': {1: 'B', 2: 'B', ...},
            '0101': {1: 'A', 2: 'B', ...},
            '0102': {1: 'D', 2: 'A', ...},
            ...
        }
    """
    wb = openpyxl.load_workbook(filepath)
    ws = wb.active
    
    # Đọc header row
    headers = []
    for cell in ws[1]:
        val = cell.value
        if val:
            headers.append(str(val).strip())
        else:
            headers.append('')
    
    print(f"  📋 Headers: {headers}")
    
    # Xác định mapping: column index → exam code
    exam_columns = {}
    for idx, header in enumerate(headers):
        if idx == 0:
            continue  # Cột "Câu"
        
        # Tìm mã đề trong header
        header_lower = header.lower()
        if 'gốc' in header_lower or 'goc' in header_lower or '000' in header_lower:
            exam_columns[idx] = 'gốc'
        else:
            # Tìm mã số: "Mã 0101" → "0101"
            code_match = re.search(r'(\d{3,4})', header)
            if code_match:
                exam_columns[idx] = code_match.group(1)
            elif header.strip():
                exam_columns[idx] = header.strip()
    
    print(f"  📋 Exam codes found: {list(exam_columns.values())}")
    
    # Đọc data rows
    answer_keys = {code: {} for code in exam_columns.values()}
    
    for row in ws.iter_rows(min_row=2, values_only=False):
        q_num_cell = row[0].value
        if q_num_cell is None:
            continue
        
        try:
            q_num = int(q_num_cell)
        except (ValueError, TypeError):
            continue
        
        for col_idx, code in exam_columns.items():
            if col_idx < len(row):
                val = row[col_idx].value
                if val:
                    answer_keys[code][q_num] = str(val).strip().upper()
    
    # Log summary
    for code, answers in answer_keys.items():
        print(f"  ✅ Mã {code}: {len(answers)} câu")
    
    return answer_keys


# ============================================================
# 3. AI QUESTION MATCHER - Khớp câu hỏi trộn với gốc
# ============================================================

def build_matching_prompt(original_questions, shuffled_questions, exam_code):
    """
    Tạo prompt cho AI để khớp câu hỏi trộn với câu gốc.
    
    AI sẽ so sánh NỘI DUNG câu hỏi (không phải đáp án, vì đáp án cũng bị trộn).
    """
    prompt = f"""Bạn là chuyên gia kiểm tra đề thi. Nhiệm vụ: khớp từng câu hỏi trong ĐỀ TRỘN (mã {exam_code}) với câu hỏi tương ứng trong ĐỀ GỐC.

LƯU Ý QUAN TRỌNG:
- Câu hỏi ĐÃ BỊ ĐẢO THỨ TỰ (câu 1 đề trộn có thể là câu 15 đề gốc)
- Các phương án A, B, C, D CŨNG ĐÃ BỊ ĐẢO THỨ TỰ
- Hãy so sánh dựa trên NỘI DUNG CÂU HỎI (phần đề bài/stem), KHÔNG dựa trên thứ tự đáp án
- Nếu câu hỏi quá ngắn hoặc tương tự nhau, hãy dùng cả nội dung đáp án để phân biệt

=== ĐỀ GỐC ===
"""
    for q in original_questions:
        opts_str = ' | '.join(f"{k}. {v}" for k, v in sorted(q.options.items()))
        prompt += f"Câu {q.number}: {q.text}\n  {opts_str}\n\n"
    
    prompt += f"\n=== ĐỀ TRỘN (Mã {exam_code}) ===\n"
    for q in shuffled_questions:
        opts_str = ' | '.join(f"{k}. {v}" for k, v in sorted(q.options.items()))
        prompt += f"Câu {q.number}: {q.text}\n  {opts_str}\n\n"
    
    prompt += """
=== YÊU CẦU ===
Với MỖI câu trong đề trộn, hãy tìm câu tương ứng trong đề gốc.

Ngoài ra, với mỗi cặp câu đã khớp, hãy xác định mapping đáp án:
- Phương án A của đề trộn tương ứng với phương án nào (A/B/C/D) của đề gốc?
- Tương tự cho B, C, D

Trả về kết quả dưới dạng JSON array, mỗi phần tử có format:
{
  "shuffled_q": <số câu đề trộn>,
  "original_q": <số câu đề gốc>,
  "option_mapping": {
    "A": "<chữ cái đáp án gốc tương ứng>",
    "B": "<chữ cái đáp án gốc tương ứng>",
    "C": "<chữ cái đáp án gốc tương ứng>",
    "D": "<chữ cái đáp án gốc tương ứng>"
  }
}

Ví dụ: Nếu câu 3 đề trộn = câu 12 đề gốc, và đáp án A đề trộn = đáp án C đề gốc:
{"shuffled_q": 3, "original_q": 12, "option_mapping": {"A": "C", "B": "A", "C": "D", "D": "B"}}

CHỈ TRẢ VỀ JSON ARRAY, KHÔNG THÊM TEXT GIẢI THÍCH NÀO KHÁC.
"""
    return prompt


def match_questions_with_ai(client, original_questions, shuffled_questions, exam_code,
                            original_file=None, shuffled_file=None, max_retries=3):
    """
    Gọi AI để khớp câu hỏi trộn với gốc.
    Có retry mechanism với exponential backoff.
    
    Nếu có file paths → gửi file gốc kèm prompt (hỗ trợ equations, hình ảnh).
    
    Returns:
        list[dict]: Kết quả matching từ AI
    """
    prompt = build_matching_prompt(original_questions, shuffled_questions, exam_code)
    
    # Chuẩn bị file paths để gửi kèm (cho equation/math support)
    file_paths = []
    if original_file and os.path.exists(original_file):
        file_paths.append(original_file)
    if shuffled_file and os.path.exists(shuffled_file):
        file_paths.append(shuffled_file)
    
    use_files = len(file_paths) > 0
    
    print(f"  🤖 Đang gọi AI để khớp mã đề {exam_code}...")
    print(f"     Prompt size: {len(prompt):,} chars")
    print(f"     Số câu gốc: {len(original_questions)}, Số câu trộn: {len(shuffled_questions)}")
    if use_files:
        print(f"     📎 Gửi kèm {len(file_paths)} file gốc (hỗ trợ equations)")
    
    # Nếu có file, thêm hướng dẫn vào prompt
    if use_files:
        prompt += "\n\nLƯU Ý: Tôi đã đính kèm file đề gốc và đề trộn. Nếu text ở trên thiếu công thức/equation, hãy tham khảo file đính kèm để so sánh chính xác hơn.\n"
    
    for attempt in range(1, max_retries + 1):
        try:
            if attempt > 1:
                wait_time = 5 * (2 ** (attempt - 2))  # 5s, 10s
                print(f"  ⏳ Retry #{attempt} sau {wait_time}s...")
                time.sleep(wait_time)
            
            response = client.send_data_to_AI(
                prompt=prompt,
                file_paths=file_paths if use_files else None,
                temperature=0.0,    # Cần kết quả chính xác nhất
                top_p=0.1,
                max_output_tokens=65535,
                response_mime_type="application/json"
            )
            
            # Parse JSON response
            cleaned = response.strip()
            if cleaned.startswith('```'):
                cleaned = re.sub(r'^```\w*\n?', '', cleaned)
                cleaned = re.sub(r'\n?```$', '', cleaned)
            
            result = json.loads(cleaned)
            
            # Validate kết quả
            if not isinstance(result, list):
                print(f"  ⚠️  AI trả về không phải list, thử lại...")
                continue
            
            # Kiểm tra mỗi entry có đủ field không
            valid_entries = []
            for entry in result:
                if isinstance(entry, dict) and 'shuffled_q' in entry and 'original_q' in entry:
                    # Đảm bảo option_mapping tồn tại
                    if 'option_mapping' not in entry:
                        entry['option_mapping'] = {}
                    valid_entries.append(entry)
            
            if len(valid_entries) < len(shuffled_questions) * 0.7:
                print(f"  ⚠️  AI chỉ trả về {len(valid_entries)}/{len(shuffled_questions)} cặp khớp, thử lại...")
                if attempt < max_retries:
                    continue
            
            print(f"  ✅ AI trả về {len(valid_entries)} cặp khớp cho mã {exam_code}")
            return valid_entries
            
        except json.JSONDecodeError as e:
            print(f"  ❌ Lỗi parse JSON (lần {attempt}): {e}")
            if attempt < max_retries:
                print(f"     Response (first 300 chars): {response[:300]}")
            else:
                print(f"  ❌ Đã thử {max_retries} lần, bỏ qua mã {exam_code}")
                return []
        except Exception as e:
            print(f"  ❌ Lỗi gọi AI (lần {attempt}): {e}")
            if attempt >= max_retries:
                raise
    
    return []


# ============================================================
# 4. ANSWER VERIFIER - Kiểm tra đáp án
# ============================================================

def verify_answers(ai_matching, answer_keys, exam_code):
    """
    Xác minh đáp án dựa trên kết quả matching từ AI.
    
    Logic:
    1. Lấy câu gốc tương ứng với câu trộn
    2. Lấy đáp án đúng của câu gốc từ bảng đáp án gốc
    3. Dùng option_mapping để tìm đáp án đúng tương ứng trong đề trộn
    4. So sánh với đáp án trong bảng đáp án trộn
    
    Returns:
        list[dict]: Danh sách lỗi, mỗi lỗi có:
            - exam_code: Mã đề
            - shuffled_q: STT câu trong đề trộn
            - current_answer: Đáp án hiện tại (sai)
            - correct_answer: Đáp án đúng
            - original_q: STT câu trong đề gốc
    """
    errors = []
    original_answers = answer_keys.get('gốc', {})
    shuffled_answers = answer_keys.get(exam_code, {})
    
    if not original_answers:
        print(f"  ⚠️  Không tìm thấy đáp án đề gốc!")
        return errors
    
    if not shuffled_answers:
        print(f"  ⚠️  Không tìm thấy đáp án mã {exam_code}!")
        return errors
    
    for match_item in ai_matching:
        shuffled_q = match_item['shuffled_q']
        original_q = match_item['original_q']
        option_mapping = match_item.get('option_mapping') or {}
        
        # Đáp án đúng trong đề gốc
        correct_original_answer = original_answers.get(original_q)
        if not correct_original_answer:
            print(f"  ⚠️  Không tìm thấy đáp án gốc cho câu {original_q}")
            continue
            
        correct_original_answer = str(correct_original_answer).strip()
        
        # Xác định loại câu hỏi
        # 1. Câu Đúng/Sai (Ví dụ: ĐSSS, SĐSĐ)
        # 2. Câu Trắc nghiệm 4 lựa chọn (A, B, C, D)
        # 3. Câu Điền từ (Ví dụ: 15.5, 12, "hàm số")
        is_true_false = len(correct_original_answer) == 4 and all(c.upper() in ['Đ', 'S'] for c in correct_original_answer)
        is_multiple_choice = len(correct_original_answer) == 1 and correct_original_answer.upper() in ['A', 'B', 'C', 'D']
        # Các đáp án còn lại là điền từ hoặc dạng khác
        
        correct_shuffled_answer = None
        
        if is_true_false:
            if option_mapping:
                # option_mapping có dạng: {'A': 'B', 'B': 'A', 'C': 'D', 'D': 'C'} (Trộn -> Gốc)
                orig_ans_dict = {
                    'A': correct_original_answer[0].upper(),
                    'B': correct_original_answer[1].upper(),
                    'C': correct_original_answer[2].upper(),
                    'D': correct_original_answer[3].upper()
                }
                
                shuffled_ans_chars = []
                for letter in ['A', 'B', 'C', 'D']:
                    orig_letter = option_mapping.get(letter, letter).upper()
                    shuffled_ans_chars.append(orig_ans_dict.get(orig_letter, 'S'))
                
                correct_shuffled_answer = "".join(shuffled_ans_chars)
            else:
                correct_shuffled_answer = correct_original_answer.upper()
                
        elif is_multiple_choice:
            for shuffled_letter, original_letter in option_mapping.items():
                if original_letter.upper() == correct_original_answer.upper():
                    correct_shuffled_answer = shuffled_letter.upper()
                    break
        else:
            # Câu điền từ (giữ nguyên đáp án, không bị trộn options)
            correct_shuffled_answer = correct_original_answer
            
        if not correct_shuffled_answer and is_multiple_choice:
            # Fallback nếu AI không trả về mapping cho câu Multiple Choice
            correct_shuffled_answer = correct_original_answer.upper()
            
        if not correct_shuffled_answer:
            print(f"  ⚠️  Không thể xác định đáp án đúng cho câu {shuffled_q} (gốc: câu {original_q}, đáp án gốc: {correct_original_answer})")
            continue
        
        # So sánh với đáp án trong bảng
        current_answer = shuffled_answers.get(shuffled_q)
        if not current_answer:
            print(f"  ⚠️  Không tìm thấy đáp án mã {exam_code} câu {shuffled_q}")
            continue
            
        current_answer = str(current_answer).strip()
        
        if current_answer.upper() != correct_shuffled_answer.upper():
            errors.append({
                'exam_code': f"Mã {exam_code}",
                'shuffled_q': shuffled_q,
                'current_answer': current_answer,
                'correct_answer': correct_shuffled_answer,
                'original_q': original_q
            })
            print(f"  ❌ SAI ĐÁP ÁN: Mã {exam_code} - Câu {shuffled_q} "
                  f"(gốc: Câu {original_q}) | "
                  f"Hiện tại: {current_answer} → Đúng: {correct_shuffled_answer}")
    
    return errors


# ============================================================
# 5. STRUCTURAL CHECKS - Kiểm tra định dạng & cấu trúc
# ============================================================

def check_duplicate_options(questions, exam_code):
    """
    Kiểm tra lỗi trùng đáp án trong một đề.
    
    Phát hiện 2 loại lỗi:
    1. Trùng nội dung: A. xyz  B. xyz (2 phương án cùng nội dung)
    2. Trùng ký hiệu: A. xyz  A. abc (cùng chữ cái xuất hiện 2 lần)
    
    Returns:
        list[dict]: Danh sách lỗi format đồng nhất
    """
    errors = []
    
    for q in questions:
        # --- Lỗi trùng nội dung đáp án ---
        opt_items = list(q.options.items())  # [('A', 'text'), ('B', 'text'), ...]
        for i in range(len(opt_items)):
            for j in range(i + 1, len(opt_items)):
                letter_i, text_i = opt_items[i]
                letter_j, text_j = opt_items[j]
                # Chuẩn hóa để so sánh (bỏ dấu cách thừa, lowercase)
                norm_i = normalize_text(text_i).lower()
                norm_j = normalize_text(text_j).lower()
                if norm_i and norm_j and norm_i == norm_j:
                    errors.append({
                        'exam_code': f"Mã {exam_code}" if exam_code != 'gốc' else 'Đề gốc',
                        'question_num': q.number,
                        'error_type': 'TRÙNG NỘI DUNG ĐÁP ÁN',
                        'detail': f'{letter_i}. và {letter_j}. cùng nội dung: "{text_i[:60]}"',
                        'original_q': q.number if exam_code == 'gốc' else None
                    })
        
        # --- Lỗi trùng ký hiệu (A xuất hiện 2 lần) ---
        # Kiểm tra trong raw text (parse lại block gốc)
        # Logic: nếu options dict có ít hơn 4 key → có thể bị trùng ký hiệu
        if len(q.options) < 4:
            missing = [l for l in 'ABCD' if l not in q.options]
            errors.append({
                'exam_code': f"Mã {exam_code}" if exam_code != 'gốc' else 'Đề gốc',
                'question_num': q.number,
                'error_type': 'THIẾU/TRÙNG KÝ HIỆU ĐÁP ÁN',
                'detail': f"Thiếu phương án: {', '.join(missing)}. Chỉ có {len(q.options)}/4 phương án.",
                'original_q': q.number if exam_code == 'gốc' else None
            })
    
    return errors


def check_duplicate_options_raw(filepath, exam_code):
    """
    Kiểm tra lỗi trùng ký hiệu đáp án ở mức raw text (A. xyz  A. abc).
    Quét trực tiếp văn bản thay vì dùng parsed questions.
    """
    errors = []
    lines = extract_all_text_from_docx(filepath)
    full_text = '\n'.join(lines)
    
    # Tìm các câu hỏi
    question_pattern = re.compile(r'(?:Câu|Question)\s*(\d+)\s*[.:]\s*', re.IGNORECASE)
    question_matches = list(question_pattern.finditer(full_text))
    
    for i, match in enumerate(question_matches):
        q_num = int(match.group(1))
        q_start = match.end()
        q_end = question_matches[i + 1].start() if i + 1 < len(question_matches) else len(full_text)
        q_block = full_text[q_start:q_end]
        
        # Đếm số lần xuất hiện của mỗi ký hiệu A., B., C., D.
        for letter in 'ABCD':
            pattern = re.compile(rf'(?:^|\s){letter}\.\s', re.MULTILINE)
            count = len(pattern.findall(q_block))
            if count > 1:
                errors.append({
                    'exam_code': f"Mã {exam_code}" if exam_code != 'gốc' else 'Đề gốc',
                    'question_num': q_num,
                    'error_type': 'TRÙNG KÝ HIỆU ĐÁP ÁN',
                    'detail': f'Ký hiệu "{letter}." xuất hiện {count} lần trong câu hỏi.',
                    'original_q': q_num if exam_code == 'gốc' else None
                })
    
    return errors


def check_layout_issues(filepath, exam_code):
    """
    Kiểm tra vỡ định dạng (Layout):
    - Các ký hiệu A, B, C, D bị nhảy Tab, không thẳng hàng
    - Phát hiện bằng cách kiểm tra XML run properties
    """
    errors = []
    doc = Document(filepath)
    
    # Duyệt từng paragraph, tìm pattern đáp án
    current_q_num = 0
    option_pattern = re.compile(r'^\s*([A-D])\.\s', re.IGNORECASE)
    question_pattern = re.compile(r'(?:Câu|Question)\s*(\d+)\s*[.:]', re.IGNORECASE)
    
    # Thu thập vị trí tab/indent của các options trong từng câu
    q_option_positions = {}  # {q_num: [(letter, indent_level, para_index)]}
    
    for para_idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        
        # Cập nhật câu hỏi hiện tại
        q_match = question_pattern.search(text)
        if q_match:
            current_q_num = int(q_match.group(1))
        
        # Kiểm tra option
        opt_match = option_pattern.match(text)
        if opt_match and current_q_num > 0:
            letter = opt_match.group(1).upper()
            
            # Lấy indent level từ paragraph format
            indent = 0
            if para.paragraph_format.left_indent:
                indent = para.paragraph_format.left_indent
            elif para.paragraph_format.first_line_indent:
                indent = para.paragraph_format.first_line_indent
            
            # Đếm tab characters ở đầu
            tab_count = 0
            for run in para.runs:
                if run.text:
                    tab_count += run.text.count('\t')
            
            if current_q_num not in q_option_positions:
                q_option_positions[current_q_num] = []
            q_option_positions[current_q_num].append({
                'letter': letter, 'indent': indent, 'tabs': tab_count, 'para_idx': para_idx
            })
    
    # So sánh indent/tab giữa các options trong cùng câu
    for q_num, positions in q_option_positions.items():
        if len(positions) < 2:
            continue
        
        # Kiểm tra tab không đồng nhất
        tab_counts = [p['tabs'] for p in positions]
        if len(set(tab_counts)) > 1:
            details = ', '.join(f"{p['letter']}: {p['tabs']} tab" for p in positions)
            errors.append({
                'exam_code': f"Mã {exam_code}" if exam_code != 'gốc' else 'Đề gốc',
                'question_num': q_num,
                'error_type': 'VỠ LAYOUT (TAB KHÔNG ĐỀU)',
                'detail': f'Số tab không đồng nhất giữa các phương án: {details}',
                'original_q': q_num if exam_code == 'gốc' else None
            })
        
        # Kiểm tra indent không đồng nhất
        indents = [p['indent'] for p in positions]
        if len(set(indents)) > 1 and any(i != 0 for i in indents):
            errors.append({
                'exam_code': f"Mã {exam_code}" if exam_code != 'gốc' else 'Đề gốc',
                'question_num': q_num,
                'error_type': 'VỠ LAYOUT (INDENT KHÔNG ĐỀU)',
                'detail': f'Căn lề không đồng nhất giữa các phương án A, B, C, D.',
                'original_q': q_num if exam_code == 'gốc' else None
            })
    
    return errors


def check_page_split_questions(filepath, exam_code):
    """
    Kiểm tra lỗi cắt đôi câu: câu hỏi nằm ở cuối trang này và đầu trang kia.
    Phát hiện bằng page break (w:lastRenderedPageBreak hoặc w:br type=page).
    """
    errors = []
    doc = Document(filepath)
    from lxml import etree
    
    nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    
    current_q_num = 0
    q_start_found = False
    question_pattern = re.compile(r'(?:Câu|Question)\s*(\d+)\s*[.:]', re.IGNORECASE)
    option_pattern = re.compile(r'^\s*[A-D]\.\s', re.IGNORECASE)
    
    # Xác định vùng của từng câu hỏi và tìm page break trong đó
    q_paras = {}  # {q_num: [para_indices]}
    para_has_page_break = {}  # {para_idx: True/False}
    
    for para_idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        # Cập nhật câu hỏi hiện tại
        q_match = question_pattern.search(text)
        if q_match:
            current_q_num = int(q_match.group(1))
        
        if current_q_num > 0:
            if current_q_num not in q_paras:
                q_paras[current_q_num] = []
            q_paras[current_q_num].append(para_idx)
        
        # Kiểm tra page break trong paragraph
        has_break = False
        para_xml = para._element
        
        # Tìm lastRenderedPageBreak (page break tự động)
        last_rendered = para_xml.findall('.//w:lastRenderedPageBreak', nsmap)
        if last_rendered:
            has_break = True
        
        # Tìm manual page break: <w:br w:type="page"/>
        breaks = para_xml.findall('.//w:br', nsmap)
        for br in breaks:
            if br.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type') == 'page':
                has_break = True
        
        para_has_page_break[para_idx] = has_break
    
    # Kiểm tra mỗi câu hỏi: nếu bên trong vùng câu hỏi có page break
    # (trừ page break ở paragraph đầu tiên của câu - đó là bình thường)
    for q_num, para_indices in q_paras.items():
        if len(para_indices) < 2:
            continue
        
        # Page break ở các paragraph GIỮA câu hỏi (không phải paragraph đầu)
        for p_idx in para_indices[1:]:
            if para_has_page_break.get(p_idx, False):
                errors.append({
                    'exam_code': f"Mã {exam_code}" if exam_code != 'gốc' else 'Đề gốc',
                    'question_num': q_num,
                    'error_type': 'CÂU BỊ CẮT ĐÔI QUA TRANG',
                    'detail': f'Câu {q_num} bị chia cắt bởi ngắt trang giữa nội dung câu hỏi.',
                    'original_q': q_num if exam_code == 'gốc' else None
                })
                break  # Chỉ báo 1 lần cho mỗi câu
    
    return errors


def check_page_count(filepath, exam_code, max_pages_normal=4, max_pages_english=6):
    """
    Kiểm tra quá số trang.
    - Môn thường: tối đa 4 trang
    - Tiếng Anh: tối đa 6 trang
    
    Đếm page break trong docx để ước lượng số trang.
    """
    errors = []
    doc = Document(filepath)
    from lxml import etree
    
    nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    
    # Đếm số page breaks
    page_break_count = 0
    body = doc.element.body
    
    # lastRenderedPageBreak
    page_break_count += len(body.findall('.//w:lastRenderedPageBreak', nsmap))
    
    # Manual page breaks
    for br in body.findall('.//w:br', nsmap):
        if br.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type') == 'page':
            page_break_count += 1
    
    # Section breaks (new page)
    for section_props in body.findall('.//w:sectPr', nsmap):
        pg_type = section_props.find('.//w:type', nsmap)
        if pg_type is not None and pg_type.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val') == 'nextPage':
            page_break_count += 1
    
    total_pages = page_break_count + 1  # Trang đầu tiên không cần break
    
    # Xác định có phải đề Tiếng Anh không (dựa trên tên file hoặc nội dung)
    basename = os.path.basename(filepath).lower()
    is_english = any(kw in basename for kw in ['ta', 'eng', 'english', 'anh'])
    
    # Cũng check nội dung nếu chưa xác định
    if not is_english:
        lines = extract_all_text_from_docx(filepath)
        first_lines = ' '.join(lines[:20]).lower()
        is_english = any(kw in first_lines for kw in ['tiếng anh', 'english', 'reading', 'mark the letter'])
    
    max_pages = max_pages_english if is_english else max_pages_normal
    subject = 'Tiếng Anh' if is_english else 'Môn thường'
    
    if total_pages > max_pages:
        errors.append({
            'exam_code': f"Mã {exam_code}" if exam_code != 'gốc' else 'Đề gốc',
            'question_num': 0,  # 0 = lỗi toàn đề
            'error_type': 'QUÁ SỐ TRANG',
            'detail': f'Đề có {total_pages} trang, vượt quy định {max_pages} trang ({subject}).',
            'original_q': None
        })
    
    print(f"  📄 Số trang: {total_pages}/{max_pages} ({subject})")
    
    return errors, total_pages


def run_structural_checks(filepath, questions, exam_code):
    """
    Chạy tất cả kiểm tra cấu trúc/định dạng cho một file đề.
    
    Returns:
        list[dict]: Tổng hợp tất cả lỗi cấu trúc
    """
    all_structural_errors = []
    
    # 1. Duplicate options (parsed)
    dup_errors = check_duplicate_options(questions, exam_code)
    if dup_errors:
        print(f"  🔴 Phát hiện {len(dup_errors)} lỗi trùng đáp án")
    all_structural_errors.extend(dup_errors)
    
    # 2. Duplicate option letters (raw text)
    dup_raw_errors = check_duplicate_options_raw(filepath, exam_code)
    if dup_raw_errors:
        print(f"  🔴 Phát hiện {len(dup_raw_errors)} lỗi trùng ký hiệu")
    all_structural_errors.extend(dup_raw_errors)
    
    # 3. Layout issues
    layout_errors = check_layout_issues(filepath, exam_code)
    if layout_errors:
        print(f"  🟡 Phát hiện {len(layout_errors)} lỗi layout")
    all_structural_errors.extend(layout_errors)
    
    # 4. Page split questions
    split_errors = check_page_split_questions(filepath, exam_code)
    if split_errors:
        print(f"  🟠 Phát hiện {len(split_errors)} câu bị cắt đôi qua trang")
    all_structural_errors.extend(split_errors)
    
    # 5. Page count
    page_errors, total_pages = check_page_count(filepath, exam_code)
    if page_errors:
        print(f"  🔴 Quá số trang cho phép!")
    all_structural_errors.extend(page_errors)
    
    return all_structural_errors


# ============================================================
# 5b. CONTENT COMPLETENESS CHECK
# ============================================================

def check_content_completeness(ai_matching, original_questions, shuffled_questions, exam_code):
    """
    Kiểm tra nội dung đề trộn có đầy đủ so với đề gốc không.
    Dùng kết quả AI matching đã có (không gọi AI thêm).
    
    Phát hiện:
    1. Câu thiếu: có trong đề gốc nhưng không match được trong đề trộn
    2. Câu thừa: có trong đề trộn nhưng không match câu nào ở đề gốc
    3. Số lượng không khớp
    
    Returns:
        list[dict]: Danh sách lỗi nội dung
    """
    errors = []
    
    # Tập hợp câu gốc được match
    matched_original_qs = set()
    matched_shuffled_qs = set()
    
    for item in ai_matching:
        matched_original_qs.add(item.get('original_q'))
        matched_shuffled_qs.add(item.get('shuffled_q'))
    
    # Lấy danh sách câu gốc
    if original_questions:
        if hasattr(original_questions[0], 'global_number'):
            # MathQuestion
            orig_q_nums = set(q.global_number for q in original_questions)
        else:
            # Question
            orig_q_nums = set(q.number for q in original_questions)
    else:
        orig_q_nums = set()
    
    # Lấy danh sách câu trộn
    if shuffled_questions:
        if hasattr(shuffled_questions[0], 'global_number'):
            shuf_q_nums = set(q.global_number for q in shuffled_questions)
        else:
            shuf_q_nums = set(q.number for q in shuffled_questions)
    else:
        shuf_q_nums = set()
    
    # Kiểm tra số lượng
    if len(orig_q_nums) != len(shuf_q_nums):
        errors.append({
            'exam_code': f"Mã {exam_code}",
            'error_type': 'SỐ LƯỢNG CÂU KHÔNG KHỚP',
            'detail': f'Đề gốc: {len(orig_q_nums)} câu, Đề trộn: {len(shuf_q_nums)} câu.',
            'missing_q': None,
        })
    
    # Câu gốc thiếu (không được match)
    missing_originals = orig_q_nums - matched_original_qs
    for q_num in sorted(missing_originals):
        errors.append({
            'exam_code': f"Mã {exam_code}",
            'error_type': 'THIẾU CÂU (so với đề gốc)',
            'detail': f'Câu {q_num} trong đề gốc không tìm thấy câu tương ứng trong đề trộn.',
            'missing_q': q_num,
        })
    
    # Câu trộn thừa (không match với câu gốc nào)
    extra_shuffled = shuf_q_nums - matched_shuffled_qs
    for q_num in sorted(extra_shuffled):
        errors.append({
            'exam_code': f"Mã {exam_code}",
            'error_type': 'CÂU THỞA (không có trong đề gốc)',
            'detail': f'Câu {q_num} trong đề trộn không match được với câu nào trong đề gốc.',
            'missing_q': q_num,
        })
    
    return errors


# ============================================================
# 6. EXCEL REPORT GENERATOR
# ============================================================

def generate_report_excel(all_errors, structural_errors, output_path, content_errors=None):
    """
    Tạo mới hoặc GHI THÊM (Append) dữ liệu vào file Excel báo cáo.
    """
    import os
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    file_exists = os.path.exists(output_path)
    if file_exists:
        print(f"  📂 Tìm thấy báo cáo cũ. Đang ghi thêm (Append) kết quả vào: {os.path.basename(output_path)}")
        wb = openpyxl.load_workbook(output_path)
    else:
        print(f"  ✨ Tạo file báo cáo mới: {os.path.basename(output_path)}")
        wb = openpyxl.Workbook()
        ws_summary = wb.active
        ws_summary.title = "Tổng hợp"
        wb.create_sheet("Sai đáp án")
        wb.create_sheet("Lỗi trùng đáp án")
        wb.create_sheet("Lỗi Layout")
        wb.create_sheet("Cắt trang - Quá trang")
        wb.create_sheet("Nội dung thừa thiếu")

    ws_summary = wb["Tổng hợp"]
    ws_sai_dap_an = wb["Sai đáp án"]
    ws_trung_dap_an = wb["Lỗi trùng đáp án"]
    ws_layout = wb["Lỗi Layout"]
    ws_loi_khac = wb["Cắt trang - Quá trang"]
    ws_content = wb["Nội dung thừa thiếu"]

    # === STYLES ===
    header_font = Font(name='Times New Roman', bold=True, size=11, color='FFFFFF')
    header_fill = PatternFill(start_color='4472C4', end_color='4472C4', fill_type='solid')
    header_alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
    data_font = Font(name='Times New Roman', size=11)
    data_alignment_center = Alignment(horizontal='center', vertical='center')
    thin_border = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'), bottom=Side(style='thin'))

    def style_header(ws, headers):
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_alignment
            cell.border = thin_border

    # 1. Lấy danh sách mã đề đang được chạy lần này
    content_errors_list = content_errors if content_errors else []
    current_exam_codes = set()
    for err in list(all_errors) + structural_errors + content_errors_list:
        if err.get('exam_code'):
            current_exam_codes.add(err['exam_code'])

    # 2. Xóa dữ liệu cũ của các mã đề này (tránh bị lặp dòng nếu chạy lại nhiều lần)
    def remove_old_rows(ws):
        if ws.max_row <= 1: return
        for row in range(ws.max_row, 1, -1):
            cell_val = ws.cell(row=row, column=1).value
            if cell_val in current_exam_codes or str(cell_val).startswith('✅'):
                ws.delete_rows(row, 1)

    if file_exists:
        for ws in [ws_sai_dap_an, ws_trung_dap_an, ws_layout, ws_loi_khac, ws_content]:
            remove_old_rows(ws)

    # Đảm bảo headers luôn tồn tại
    headers_sai = ['Mã đề', 'STT câu (đề trộn)', 'Đáp án hiện tại', 'Đáp án đúng', 'Câu gốc tương ứng']
    style_header(ws_sai_dap_an, headers_sai)
    headers_struct = ['Mã đề', 'STT câu', 'Loại lỗi', 'Chi tiết lỗi', 'Câu gốc tương ứng']
    style_header(ws_trung_dap_an, headers_struct)
    style_header(ws_layout, headers_struct)
    style_header(ws_loi_khac, headers_struct)
    headers_content = ['Mã đề', 'Loại lỗi', 'Chi tiết', 'Câu liên quan']
    style_header(ws_content, headers_content)

    # 3. Hàm phụ trợ ghi Data nối tiếp
    def append_data(ws, data_list, is_sai=False, is_content=False):
        for err in data_list:
            r_idx = ws.max_row + 1
            ws.cell(row=r_idx, column=1, value=err.get('exam_code', '')).alignment = data_alignment_center
            if is_sai:
                ws.cell(row=r_idx, column=2, value=err.get('shuffled_q', '')).alignment = data_alignment_center
                ws.cell(row=r_idx, column=3, value=err.get('current_answer', '')).alignment = data_alignment_center
                ws.cell(row=r_idx, column=3).font = Font(name='Times New Roman', size=11, bold=True, color='FF0000')
                ws.cell(row=r_idx, column=4, value=err.get('correct_answer', '')).alignment = data_alignment_center
                ws.cell(row=r_idx, column=4).font = Font(name='Times New Roman', size=11, bold=True, color='008000')
                ws.cell(row=r_idx, column=5, value=err.get('original_q', '')).alignment = data_alignment_center
            elif is_content:
                ws.cell(row=r_idx, column=2, value=err.get('error_type', '')).alignment = data_alignment_center
                ws.cell(row=r_idx, column=3, value=err.get('detail', '')).alignment = Alignment(horizontal='left', vertical='center', wrap_text=True)
                ws.cell(row=r_idx, column=4, value=err.get('missing_q', '')).alignment = data_alignment_center
            else:
                q_val = err.get('question_num', 0)
                ws.cell(row=r_idx, column=2, value=q_val if q_val > 0 else 'Toàn đề').alignment = data_alignment_center
                ws.cell(row=r_idx, column=3, value=err.get('error_type', '')).alignment = data_alignment_center
                ws.cell(row=r_idx, column=4, value=err.get('detail', '')).alignment = Alignment(horizontal='left', vertical='center', wrap_text=True)
                ws.cell(row=r_idx, column=5, value=err.get('original_q', '')).alignment = data_alignment_center
            
            # Kẻ khung
            cols_count = 5 if not is_content else 4
            for c in range(1, cols_count + 1):
                ws.cell(row=r_idx, column=c).border = thin_border
                if not (is_sai and c in [3, 4]): ws.cell(row=r_idx, column=c).font = data_font

    # Phân loại lỗi cấu trúc
    trung_errors = [e for e in structural_errors if 'TRÙNG' in e.get('error_type', '')]
    layout_errors = [e for e in structural_errors if 'LAYOUT' in e.get('error_type', '')]
    other_errors = [e for e in structural_errors if 'TRÙNG' not in e.get('error_type', '') and 'LAYOUT' not in e.get('error_type', '')]

    # Ghi nối dữ liệu
    append_data(ws_sai_dap_an, sorted(all_errors, key=lambda x: (x['exam_code'], x.get('shuffled_q', 0))), is_sai=True)
    append_data(ws_trung_dap_an, sorted(trung_errors, key=lambda x: (x.get('exam_code', ''), x.get('question_num', 0))))
    append_data(ws_layout, sorted(layout_errors, key=lambda x: (x.get('exam_code', ''), x.get('question_num', 0))))
    append_data(ws_loi_khac, sorted(other_errors, key=lambda x: (x.get('exam_code', ''), x.get('question_num', 0))))
    append_data(ws_content, sorted(content_errors_list, key=lambda x: (x.get('exam_code', ''), x.get('missing_q', 0))), is_content=True)

    # Căn chỉnh độ rộng cột
    for ws in [ws_sai_dap_an, ws_trung_dap_an, ws_layout, ws_loi_khac]:
        ws.column_dimensions['A'].width = 15; ws.column_dimensions['B'].width = 18
        ws.column_dimensions['C'].width = 30 if ws != ws_sai_dap_an else 18
        ws.column_dimensions['D'].width = 50 if ws != ws_sai_dap_an else 18
        ws.column_dimensions['E'].width = 20
    ws_content.column_dimensions['A'].width = 15; ws_content.column_dimensions['B'].width = 30
    ws_content.column_dimensions['C'].width = 50; ws_content.column_dimensions['D'].width = 15

    # Thêm câu "Không có lỗi" nếu sheet trống trơn
    def check_empty(ws, msg, cols):
        if ws.max_row == 1:
            ws.cell(row=2, column=1, value=msg).font = Font(name='Times New Roman', size=11, color='008000')
            ws.merge_cells(f'A2:{cols}2')
    check_empty(ws_sai_dap_an, "✅ Không có bài nào sai đáp án", "E")
    check_empty(ws_trung_dap_an, "✅ Không có lỗi trùng đáp án", "E")
    check_empty(ws_layout, "✅ Không có lỗi layout", "E")
    check_empty(ws_loi_khac, "✅ Không có lỗi in ấn/cắt trang", "E")
    check_empty(ws_content, "✅ Nội dung đầy đủ so với đề gốc", "D")

    # 4. TÍNH TOÁN LẠI SHEET TỔNG HỢP (Dựa trên tất cả dữ liệu có trong file)
    ws_summary.delete_rows(1, ws_summary.max_row)
    ws_summary.cell(row=1, column=1, value="TỔNG HỢP KẾT QUẢ KIỂM TRA ĐỀ").font = Font(name='Times New Roman', size=14, bold=True, color='4472C4')
    ws_summary.merge_cells('A1:D1')

    def count_errors(ws):
        return sum(1 for row in range(2, ws.max_row + 1) if ws.cell(row=row, column=1).value and not str(ws.cell(row=row, column=1).value).startswith('✅'))

    ws_summary.cell(row=3, column=1, value="Lỗi sai đáp án:").font = Font(name='Times New Roman', size=11, bold=True)
    ws_summary.cell(row=3, column=2, value=count_errors(ws_sai_dap_an)).font = Font(name='Times New Roman', size=11, bold=True, color='FF0000')
    ws_summary.cell(row=4, column=1, value="Lỗi trùng đáp án:").font = Font(name='Times New Roman', size=11, bold=True)
    ws_summary.cell(row=4, column=2, value=count_errors(ws_trung_dap_an)).font = Font(name='Times New Roman', size=11, bold=True, color='ED7D31')
    ws_summary.cell(row=5, column=1, value="Lỗi vỡ Layout:").font = Font(name='Times New Roman', size=11, bold=True)
    ws_summary.cell(row=5, column=2, value=count_errors(ws_layout)).font = Font(name='Times New Roman', size=11, bold=True, color='ED7D31')
    ws_summary.cell(row=6, column=1, value="Lỗi cắt trang/quá trang:").font = Font(name='Times New Roman', size=11, bold=True)
    ws_summary.cell(row=6, column=2, value=count_errors(ws_loi_khac)).font = Font(name='Times New Roman', size=11, bold=True, color='ED7D31')
    ws_summary.cell(row=7, column=1, value="Nội dung thừa/thiếu:").font = Font(name='Times New Roman', size=11, bold=True)
    ws_summary.cell(row=7, column=2, value=count_errors(ws_content)).font = Font(name='Times New Roman', size=11, bold=True, color='ED7D31')

    # Thống kê chi tiết từng mã đề có trong Excel
    error_by_code = {}
    for ws in [ws_sai_dap_an, ws_trung_dap_an, ws_layout, ws_loi_khac, ws_content]:
        for row in range(2, ws.max_row + 1):
            code = ws.cell(row=row, column=1).value
            if code and not str(code).startswith('✅'):
                error_by_code[code] = error_by_code.get(code, 0) + 1
    
    # Đảm bảo các mã đề vừa chạy nếu có 0 lỗi thì vẫn hiện diện
    for code in current_exam_codes:
        if code not in error_by_code: error_by_code[code] = 0

    ws_summary.cell(row=9, column=1, value="Mã đề").font = Font(name='Times New Roman', size=11, bold=True)
    ws_summary.cell(row=9, column=2, value="Tổng số lỗi").font = Font(name='Times New Roman', size=11, bold=True)
    r = 10
    for code, count in sorted(error_by_code.items()):
        ws_summary.cell(row=r, column=1, value=code).font = data_font
        ws_summary.cell(row=r, column=2, value=count).font = Font(name='Times New Roman', size=11, color='FF0000' if count > 0 else '008000')
        r += 1

    ws_summary.column_dimensions['A'].width = 25; ws_summary.column_dimensions['B'].width = 15

    wb.save(output_path)
    print(f"\n📊 Báo cáo đã được cập nhật thành công tại: {output_path}")
# ============================================================
# 6. MAIN - Orchestrator
# ============================================================

def find_exam_files(base_dir):
    """
    Tự động tìm các file đề trong thư mục.
    
    Returns:
        tuple: (original_file, answer_key_file, list_of_shuffled_files)
    """
    original_file = None
    answer_key_file = None
    shuffled_files = []
    
    # Tìm trong thư mục chính và các thư mục con
    search_dirs = [base_dir]
    for item in os.listdir(base_dir):
        item_path = os.path.join(base_dir, item)
        if os.path.isdir(item_path) and item != 'venv' and not item.startswith('.'):
            search_dirs.append(item_path)
    
    for search_dir in search_dirs:
        for file in os.listdir(search_dir):
            filepath = os.path.join(search_dir, file)
            if not os.path.isfile(filepath):
                continue
            
            file_lower = file.lower()
            
            # File đáp án Excel
            if file_lower.endswith('.xlsx') and ('đáp án' in file_lower or 'dap an' in file_lower or 'answer' in file_lower or 'ĐA' in file_lower):
                answer_key_file = filepath
            
            # File đề gốc (hỗ trợ cả docx và pdf)
            elif (file_lower.endswith('.docx') or file_lower.endswith('.pdf')) and ('gốc' in file_lower or 'goc' in file_lower or 'master' in file_lower):
                original_file = filepath
            
            # File đề trộn (có mã đề dạng số) - hỗ trợ cả docx và pdf
            elif (file_lower.endswith('.docx') or file_lower.endswith('.pdf')) and re.search(r'_\d{4}\.(docx|pdf)$', file_lower):
                shuffled_files.append(filepath)
    
    # Sắp xếp file trộn theo mã đề
    shuffled_files.sort()
    
    return original_file, answer_key_file, shuffled_files


def extract_exam_code(filepath):
    """Trích xuất mã đề từ tên file, ví dụ Ma_de_0101.pdf → 0101, PHC12_0101.docx → 0101"""
    basename = os.path.basename(filepath)
    # Thử pattern: Ma_de_XXXX hoặc _XXXX
    match = re.search(r'(?:Ma_de_|_)(\d{3,4})', basename, re.IGNORECASE)
    if match:
        return match.group(1)
    # Thử tìm số 4 chữ số bất kỳ trong tên file
    match = re.search(r'(\d{4})', basename)
    if match:
        return match.group(1)
    return os.path.splitext(basename)[0]


def _extract_text_from_file(filepath):
    """Helper: trích xuất text từ file (docx hoặc pdf)."""
    ext = os.path.splitext(filepath)[1].lower()
    if ext == '.docx':
        return extract_all_text_from_docx(filepath)
    elif ext == '.pdf':
        return extract_text_from_pdf(filepath)
    return []


def main():
    parser = argparse.ArgumentParser(description="Công cụ kiểm tra đáp án đề trộn (Check Đề)")
    parser.add_argument('-o', '--original', type=str, help='Đường dẫn file đề gốc (.docx hoặc .pdf)')
    parser.add_argument('-a', '--answer', type=str, help='Đường dẫn file đáp án (.xlsx)')
    parser.add_argument('-s', '--shuffled', type=str, nargs='+', help='Danh sách file đề trộn (.docx, .pdf) hoặc thư mục')
    
    # THÊM ĐÚNG 1 DÒNG NÀY CHO FILE LẺ:
    parser.add_argument('-f', '--file', type=str, nargs='+', help='Chỉ định đích danh 1 (hoặc vài) file đề trộn cụ thể')
    parser.add_argument('--vision', action='store_true', help='Bắt buộc dùng vision mode')
    # THÊM DÒNG NÀY:
    parser.add_argument('--subject', type=str, default='auto', help='Ép kiểu môn học (math, english, ...)')
    
    args = parser.parse_args()
    
    print("=" * 60)
    print("  CHECK ĐỀ - Công cụ kiểm tra đáp án đề trộn")
    print("=" * 60)
    
    # SỬA LỖI: Lấy đúng thư mục chứa file .exe thay vì thư mục Temp của PyInstaller
    if getattr(sys, 'frozen', False):
        base_dir = os.path.dirname(sys.executable)
    else:
        base_dir = os.path.dirname(os.path.abspath(__file__))
    
    original_file = args.original
    answer_key_file = args.answer
    shuffled_files = []
    force_vision = getattr(args, 'vision', False)
    
    if args.shuffled:
        for path in args.shuffled:
            if os.path.isfile(path):
                shuffled_files.append(os.path.abspath(path))
            elif os.path.isdir(path):
                for f in os.listdir(path):
                    fpath = os.path.join(path, f)
                    if os.path.isfile(fpath) and fpath.lower().endswith(('.docx', '.pdf')):
                        shuffled_files.append(os.path.abspath(fpath))
    if getattr(args, 'file', None):
        for path in args.file:
            if os.path.isfile(path) and path not in shuffled_files:
                shuffled_files.append(os.path.abspath(path))
    # --- STEP 1: Tìm file ---
    print("\n📁 BƯỚC 1: Tìm kiếm/Tải file...")
    if not original_file or not answer_key_file or not shuffled_files:
        print("   (Đang tìm tự động cho các thông số chưa được chỉ định...)")
        auto_orig, auto_ans, auto_shuf = find_exam_files(base_dir)
        if not original_file: original_file = auto_orig
        if not answer_key_file: answer_key_file = auto_ans
        if not shuffled_files: shuffled_files = auto_shuf
    
    if not original_file:
        print("❌ Không tìm thấy file đề gốc! Dùng tham số -o hoặc đảm bảo tên file có chữ 'gốc'/'master'")
        return
    if not answer_key_file:
        print("❌ Không tìm thấy file đáp án Excel! Dùng tham số -a hoặc đảm bảo tên file có chữ 'đáp án'/'ĐA'")
        return
    if not shuffled_files:
        print("❌ Không tìm thấy file đề trộn! Dùng tham số -s hoặc để chung thư mục")
        return
    
    original_file = os.path.abspath(original_file)
    answer_key_file = os.path.abspath(answer_key_file)
    orig_ext = os.path.splitext(original_file)[1].lower()
    
    print(f"  📄 Đề gốc: {os.path.basename(original_file)} ({orig_ext})")
    print(f"  📊 Đáp án: {os.path.basename(answer_key_file)}")
    print(f"  📝 Đề trộn: {len(shuffled_files)} file")
    for f in shuffled_files:
        print(f"     - {os.path.basename(f)}")
    
    # --- STEP 2: Detect subject ---
    print("\n🎯 BƯỚC 2: Nhận diện môn học...")
    orig_lines = _extract_text_from_file(original_file)
    
    if args.subject and args.subject != 'auto':
        subject = args.subject
    else:
        if MATH_HANDLER_AVAILABLE:
            subject = detect_subject(original_file, orig_lines if orig_lines else None)
        else:
            subject = 'other'
    
    subject_display = get_subject_label(subject) if PROMPT_LOADER_AVAILABLE else subject.upper()
    print(f"  🎯 Môn: {subject_display} ({subject})")
    
    # Quyết định dùng math handler hay flow cũ
    use_math_handler = MATH_HANDLER_AVAILABLE and subject in ('math',)
    # Vision mode: bắt buộc cho Toán/KHTN có công thức, hoặc khi đề gốc là PDF
    use_vision = force_vision or subject in ('math',) or orig_ext == '.pdf'
    if use_vision:
        print(f"  👁️  Vision mode: BẬT (gửi file trực tiếp cho AI)")
    
    # --- STEP 2b: Parse đề gốc (cho structural check và content comparison) ---
    print("\n📖 BƯỚC 2b: Parse đề gốc...")
    if use_math_handler and subject == 'math' and orig_lines:
        from math_exam_handler import parse_math_exam_from_text
        original_questions = parse_math_exam_from_text(
            '\n'.join(orig_lines), os.path.basename(original_file), subject=subject
        )
        print(f"  ✅ Đã parse {len(original_questions)} câu hỏi từ đề gốc (math handler)")
    elif orig_ext == '.docx':
        original_questions = parse_docx_questions(original_file)
        print(f"  ✅ Đã parse {len(original_questions)} câu hỏi từ đề gốc (docx)")
    elif orig_ext == '.pdf' and orig_lines:
        original_questions = parse_questions_from_text('\n'.join(orig_lines), os.path.basename(original_file))
        print(f"  ✅ Đã parse {len(original_questions)} câu hỏi từ đề gốc (pdf)")
    else:
        original_questions = []
        print(f"  ⚠️  Không parse được đề gốc (sẽ dùng vision mode + AI)")
    
    # --- STEP 2c: Kiểm tra cấu trúc đề gốc ---
    print("\n🔍 BƯỚC 2c: Kiểm tra cấu trúc đề gốc...")
    all_structural_errors = []
    if orig_ext == '.docx' and original_questions:
        struct_errors_goc = run_structural_checks(original_file, original_questions, 'gốc')
        all_structural_errors.extend(struct_errors_goc)
        if struct_errors_goc:
            print(f"  ⚠️  Đề gốc: {len(struct_errors_goc)} lỗi cấu trúc")
        else:
            print(f"  ✅ Đề gốc: Không có lỗi cấu trúc")
    else:
        print(f"  ℹ️  Bỏ qua kiểm tra cấu trúc đề gốc (file {orig_ext})")
    
    # --- STEP 3: Parse bảng đáp án ---
    print("\n📊 BƯỚC 3: Đọc bảng đáp án Excel...")
    answer_keys = parse_answer_key_excel(answer_key_file)
    
    # --- STEP 4: Khởi tạo AI Client ---
    print("\n🤖 BƯỚC 4: Khởi tạo AI Client...")
    project_id = os.getenv("PROJECT_ID", "onluyen-media")
    model_name = "gemini-3.1-pro-preview"
    
    creds = get_vertex_ai_credentials()
    if not creds:
        print("❌ Không thể tạo credentials!")
        return
    
    client = VertexClient(
        project_id=project_id,
        creds=creds,
        model_name=model_name,
        region="global"
    )
    
    if not client.client:
        print("❌ Không thể khởi tạo AI Client!")
        return
    
    # --- STEP 5: Kiểm tra từng mã đề ---
    print("\n🔍 BƯỚC 5: Kiểm tra từng mã đề...")
    all_errors = []
    all_content_errors = []
    debug_data_list = []
    
    for shuffled_file in shuffled_files:
        exam_code = extract_exam_code(shuffled_file)
        file_ext = os.path.splitext(shuffled_file)[1].lower()
        print(f"\n{'─' * 50}")
        print(f"  🔎 Đang kiểm tra mã đề: {exam_code}")
        print(f"     File: {os.path.basename(shuffled_file)} ({file_ext})")
        
        # --- 5a: Kiểm tra cấu trúc đề trộn ---
        struct_errors = []
        shuffled_questions = []
        
        if file_ext == '.docx':
            if not use_math_handler:
                shuffled_questions = parse_exam_file(shuffled_file)
                struct_errors = run_structural_checks(shuffled_file, shuffled_questions, exam_code)
        elif file_ext == '.pdf':
            if not use_math_handler:
                shuf_lines = extract_text_from_pdf(shuffled_file)
                if shuf_lines:
                    shuffled_questions = parse_questions_from_text('\n'.join(shuf_lines), os.path.basename(shuffled_file))
                struct_errors = check_duplicate_options(shuffled_questions, exam_code) if shuffled_questions else []
            
            if PDF_LIB:
                page_count = get_pdf_page_count(shuffled_file)
                max_pages = load_max_pages(subject) if PROMPT_LOADER_AVAILABLE else (6 if subject == 'english' else 4)
                if page_count > max_pages:
                    struct_errors.append({
                        'exam_code': f"Mã {exam_code}",
                        'question_num': 0,
                        'error_type': 'QUÁ SỐ TRANG',
                        'detail': f'Đề có {page_count} trang, vượt quy định {max_pages} trang.',
                        'original_q': None
                    })
                print(f"  📄 Số trang PDF: {page_count}/{max_pages}")
        
        if shuffled_questions:
            print(f"  📖 Đã parse {len(shuffled_questions)} câu từ đề trộn {exam_code}")
        all_structural_errors.extend(struct_errors)
        
        # --- 5b: Khớp câu hỏi bằng AI ---
        ai_matching = []
        errors = []
        
        try:
            if use_math_handler:
                # Dùng math handler v2 (hỗ trợ vision mode + 3 loại câu)
                errors, ai_matching, debug_info, orig_qs_math, shuf_qs_math = process_math_exam_v2(
                    original_file=original_file,
                    shuffled_file=shuffled_file,
                    answer_keys_raw=answer_keys,
                    exam_code=exam_code,
                    ai_client=client,
                    extract_text_fn=_extract_text_from_file,
                    use_vision=use_vision,
                )
                
                # Cập nhật kết quả parse cho 5c
                original_questions = orig_qs_math
                shuffled_questions = shuf_qs_math
                
                # Thực hiện structural checks đặc thù cho Toán
                if shuf_qs_math:
                    # Chỉ check A B C D cho PHẦN 1
                    mc_only = [q for q in shuf_qs_math if getattr(q, 'part', 1) == 1]
                    struct_errors.extend(check_duplicate_options(mc_only, exam_code))
                
                if file_ext == '.docx':
                    struct_errors.extend(check_layout_issues(shuffled_file, exam_code))
                    struct_errors.extend(check_page_split_questions(shuffled_file, exam_code))
            else:
                # Flow cũ cho Tiếng Anh / môn khác
                if use_vision and PROMPT_LOADER_AVAILABLE:
                    # Vision mode: gửi file trực tiếp
                    prompt = load_vision_prompt(subject, exam_code)
                    if not prompt:
                        prompt = build_matching_prompt(original_questions, shuffled_questions, exam_code)
                    
                    from math_exam_handler import prepare_files_for_ai
                    ai_files = prepare_files_for_ai(original_file, shuffled_file)
                    
                    print(f"  👁️  Vision mode: gửi file trực tiếp cho AI")
                    response = client.send_data_to_AI(
                        prompt=prompt,
                        file_paths=ai_files,
                        temperature=0.0,
                        top_p=0.1,
                        max_output_tokens=65535,
                        response_mime_type="application/json"
                    )
                    cleaned = response.strip()
                    if cleaned.startswith('```'):
                        cleaned = re.sub(r'^```\w*\n?', '', cleaned)
                        cleaned = re.sub(r'\n?```$', '', cleaned)
                    ai_matching = json.loads(cleaned)
                    if not isinstance(ai_matching, list):
                        ai_matching = []
                    print(f"  ✅ AI trả về {len(ai_matching)} cặp khớp")
                    errors = verify_answers(ai_matching, answer_keys, exam_code)
                else:
                    ai_matching = match_questions_with_ai(
                        client, original_questions, shuffled_questions, exam_code,
                        original_file=original_file, shuffled_file=shuffled_file
                    )
                    if ai_matching:
                        errors = verify_answers(ai_matching, answer_keys, exam_code)
                
                debug_info = {
                    'exam_code': exam_code,
                    'subject': subject,
                    'ai_matches': len(ai_matching),
                    'vision_mode': use_vision,
                }
        except Exception as e:
            print(f"  ❌ Lỗi khi gọi AI cho mã {exam_code}: {e}")
            traceback.print_exc()
            continue
        
        all_errors.extend(errors)
        
        # --- 5c: Kiểm tra nội dung thừa/thiếu ---
        if ai_matching and original_questions:
            content_errs = check_content_completeness(
                ai_matching, original_questions, shuffled_questions, exam_code
            )
            all_content_errors.extend(content_errs)
            if content_errs:
                print(f"  🟡 Mã {exam_code}: {len(content_errs)} lỗi nội dung thừa/thiếu")
        
        # Cập nhật original_q cho structural errors
        if ai_matching:
            mapping = {m['shuffled_q']: m['original_q'] for m in ai_matching}
            for serr in struct_errors:
                if serr.get('original_q') is None and serr.get('question_num', 0) > 0:
                    serr['original_q'] = mapping.get(serr['question_num'])
        
        # Lưu debug data
        debug_data_list.append({
            'exam_code': exam_code,
            'subject': subject,
            'num_shuffled_questions': len(shuffled_questions),
            'num_matches': len(ai_matching),
            'matching_results': ai_matching,
            'answer_errors': errors,
            'structural_errors': struct_errors,
        })
        
        if not errors:
            print(f"  ✅ Mã {exam_code}: Tất cả đáp án đúng!")
        else:
            print(f"  ⚠️  Mã {exam_code}: Phát hiện {len(errors)} lỗi đáp án")
        
        if struct_errors:
            print(f"  🔶 Mã {exam_code}: {len(struct_errors)} lỗi cấu trúc/định dạng")
        
        # Rate limiting
        time.sleep(2)
    
    # Lưu debug data
    debug_path = os.path.join(base_dir, "debug_check_de.json")
    try:
        with open(debug_path, 'w', encoding='utf-8') as f:
            json.dump({
                'subject': subject,
                'original_file': os.path.basename(original_file),
                'exam_results': debug_data_list,
                'total_answer_errors': len(all_errors),
                'total_structural_errors': len(all_structural_errors),
                'total_content_errors': len(all_content_errors),
            }, f, ensure_ascii=False, indent=2)
        print(f"\n  📝 Debug data đã lưu tại: {debug_path}")
    except Exception as e:
        print(f"  ⚠️  Không thể lưu debug data: {e}")
    
    print(f"\n{'=' * 60}")
    print(f"📊 BƯỚC 6: Xuất báo cáo...")
    
    # Tên file output bao gồm môn học + tên đề gốc + thời gian → KHÔNG BỊ GHI ĐÈ
    subject_name_map = {'math': 'Toan', 'english': 'TiengAnh', 'other': 'MonKhac'}
    
    # Ưu tiên lấy subject (mã môn học) trực tiếp, nếu không có thì fallback về MonKhac
    subject_label = subject_name_map.get(subject, subject.capitalize())
    
    orig_basename = os.path.splitext(os.path.basename(original_file))[0]
    safe_name = re.sub(r'[^\w\s-]', '', orig_basename).strip().replace(' ', '_')
    
    # Lấy thời gian hiện tại theo định dạng: NămThángNgày_GiờPhútGiây (VD: 20260328_174500)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    
    # Ghép timestamp vào tên file
    output_filename = f"KetQua_KiemTra_{subject_label}_{safe_name}_{timestamp}.xlsx"
    output_path = os.path.join(base_dir, output_filename)
    
    generate_report_excel(all_errors, all_structural_errors, output_path, content_errors=all_content_errors)
    
    # Tổng kết
    print(f"\n{'=' * 60}")
    print(f"  📋 TỔNG KẾT")
    print(f"  Môn học: {subject_display}")
    print(f"  Tổng số mã đề đã kiểm tra: {len(shuffled_files)}")
    print(f"  Tổng số lỗi đáp án: {len(all_errors)}")
    print(f"  Tổng số lỗi cấu trúc/định dạng: {len(all_structural_errors)}")
    print(f"  Tổng số lỗi nội dung thừa/thiếu: {len(all_content_errors)}")
    
    if all_errors:
        print(f"\n  📌 Chi tiết lỗi đáp án:")
        for err in all_errors:
            print(f"    • {err['exam_code']} - Câu {err['shuffled_q']}: "
                  f"{err['current_answer']} → {err['correct_answer']} "
                  f"(Gốc: Câu {err['original_q']})")
    
    if all_structural_errors:
        print(f"\n  📌 Chi tiết lỗi cấu trúc:")
        for err in all_structural_errors:
            q_display = f"Câu {err['question_num']}" if err.get('question_num', 0) > 0 else "Toàn đề"
            print(f"    • {err['exam_code']} - {q_display}: "
                  f"[{err['error_type']}] {err['detail']}")
    
    if all_content_errors:
        print(f"\n  📌 Chi tiết lỗi nội dung:")
        for err in all_content_errors:
            print(f"    • {err['exam_code']}: [{err['error_type']}] {err['detail']}")
    
    if not all_errors and not all_structural_errors and not all_content_errors:
        print(f"\n  ✅ KHÔNG PHÁT HIỆN LỖI NÀO!")
    
    print(f"{'=' * 60}")


if __name__ == '__main__':
    main()