import os
import sys
from dotenv import load_dotenv
from google.oauth2 import service_account
from google import genai
from google.genai import types

if hasattr(sys.stdout, 'encoding') and sys.stdout.encoding:
    if sys.stdout.encoding.lower() != 'utf-8':
        try:
            sys.stdout.reconfigure(encoding='utf-8')
        except AttributeError:
            pass
# --- LOGIC TÌM ENV ĐA NĂNG ---
# 1. Xác định vị trí file này (modules/common)

sys.path.append(os.path.dirname(__file__))
env_path = os.path.join(os.path.dirname(__file__), ".env")
# 3. Load file .env.gen
print(f"[API] Loading config")
if os.path.exists(env_path):
    load_dotenv(env_path, override=True)
    print("[API] Loaded successfully")
else:
    print(f"❌ [API] CẢNH BÁO: Không tìm thấy file tại {env_path}")
 
# ============================================================
# 2. HÀM TẠO CREDENTIALS (PUBLIC HELPER)
# ============================================================
def get_vertex_ai_credentials():
    """
    Hàm helper để lấy credentials, dùng chung cho cả callAPI và text2Image.
    """
    try:
        private_key = os.getenv("PRIVATE_KEY")
        if not private_key:
            print("❌ [API] Lỗi: Không tìm thấy PRIVATE_KEY trong .env")
            return None

        service_account_data = {
            "type": os.getenv("TYPE"),
            "project_id": os.getenv("PROJECT_ID"),
            "private_key_id": os.getenv("PRIVATE_KEY_ID"),
            "private_key": private_key.replace('\\n', '\n'),
            "client_email": os.getenv("CLIENT_EMAIL"),
            "client_id": os.getenv("CLIENT_ID"),
            "auth_uri": os.getenv("AUTH_URI"),
            "token_uri": os.getenv("TOKEN_URI"),
            "auth_provider_x509_cert_url": os.getenv("AUTH_PROVIDER_X509_CERT_URL"),
            "client_x509_cert_url": os.getenv("CLIENT_X509_CERT_URL"),
            "universe_domain": os.getenv("UNIVERSE_DOMAIN")
        }
        
        creds = service_account.Credentials.from_service_account_info(
            service_account_data,
            scopes=["https://www.googleapis.com/auth/cloud-platform"]
        )
        return creds
    except Exception as e:
        print(f"❌ [API] Lỗi khi tạo credentials: {e}")
        return None

# ============================================================
# 3. CLASS VERTEX CLIENT (CHO TEXT GENERATION)
# ============================================================

class VertexClient:
    def __init__(self, project_id, creds, model_name, region="global"):
        """
        Khởi tạo Client sử dụng google.genai SDK mới
        """
        self.model_name = model_name
        # Cache: local_path → uploaded File object (tồn tại trong session)
        # None = đã thử upload nhưng File API không khả dụng
        self._uploaded_files: dict = {}
        # False sau lần đầu xác nhận Vertex AI không hỗ trợ File API
        self._file_api_available: bool = True

        if not creds:
            print("❌ Lỗi: Credentials bị None.")
            return

        try:
            # Khởi tạo Client theo chuẩn mới
            self.client = genai.Client(
                vertexai=True,
                project=project_id,
                location=region,
                credentials=creds
            )
            print(f"✅ Init GenAI Client thành công với model: {self.model_name}")
        except Exception as e:
            print(f"Lỗi init GenAI Client: {e}")
            self.client = None

    def upload_files_cached(self, file_paths: list) -> list:
        """
        Upload các file lên Gemini File API, cache kết quả trong session.
        Nếu Vertex AI không hỗ trợ File API, tự động fallback inline toàn bộ.
        """
        # Nếu đã biết File API không khả dụng → fallback ngay, không thử nữa
        if not self._file_api_available:
            return self._inline_parts(file_paths)

        parts = []
        for path in file_paths:
            if not os.path.exists(path):
                print(f"  ⚠️  Bỏ qua file không tồn tại: {path}")
                continue

            label = os.path.basename(path)

            # Kiểm tra cache session (None = đã thử, không khả dụng)
            if path in self._uploaded_files:
                cached = self._uploaded_files[path]
                if cached is None:
                    parts.extend(self._inline_parts([path]))
                    continue
                print(f"  ⚡ File API cache hit: {label}")
                parts.append(types.Part.from_uri(
                    file_uri=cached.uri,
                    mime_type=cached.mime_type,
                ))
                continue

            # Upload mới
            try:
                lower_path = path.lower()
                if lower_path.endswith(".docx"):
                    mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                elif lower_path.endswith(".md"):
                    mime = "text/plain"
                elif lower_path.endswith(".pdf"):
                    mime = "application/pdf"
                else:
                    mime = "application/octet-stream"
                print(f"  ☁️  Đang upload {label} lên File API...")
                uploaded = self.client.files.upload(
                    file=path,
                    config=types.UploadFileConfig(mime_type=mime, display_name=label),
                )
                self._uploaded_files[path] = uploaded
                print(f"  ✅ Upload xong: {label} → {uploaded.uri}")
                parts.append(types.Part.from_uri(
                    file_uri=uploaded.uri,
                    mime_type=uploaded.mime_type,
                ))
            except Exception as e:
                err = str(e)
                is_unsupported = any(phrase in err.lower() for phrase in [
                    "only supported in the gemini developer",
                    "not supported",
                    "vertexai",
                    "404",
                    "method not found",
                ])
                if is_unsupported:
                    print(f"  ⚠️  File API không khả dụng trên Vertex AI — chuyển sang inline cho toàn session")
                    self._file_api_available = False
                    self._uploaded_files[path] = None
                    # Fallback tất cả file còn lại trong danh sách
                    remaining_idx = file_paths.index(path)
                    return parts + self._inline_parts(file_paths[remaining_idx:])
                else:
                    print(f"  ❌ Upload thất bại {label}: {e}")
                    self._uploaded_files[path] = None
                    parts.extend(self._inline_parts([path]))
        return parts

    def _inline_parts(self, file_paths: list) -> list:
        """Đọc file và trả về Parts inline (text hoặc bytes) — fallback khi File API không dùng được."""
        parts = []
        for file_path in file_paths:
            label = os.path.basename(file_path)
            try:
                if file_path.lower().endswith(".md"):
                    with open(file_path, "r", encoding="utf-8") as f:
                        md_text = f.read()
                    parts.append(types.Part.from_text(
                        text=f"\n--- BẮT ĐẦU TÀI LIỆU: {label} ---\n{md_text}\n--- KẾT THÚC TÀI LIỆU: {label} ---\n"
                    ))
                elif file_path.lower().endswith(".pdf"):
                    with open(file_path, "rb") as f:
                        pdf_bytes = f.read()
                    parts.append(types.Part.from_bytes(data=pdf_bytes, mime_type="application/pdf"))
                elif file_path.lower().endswith(".docx"):
                    # Vertex AI không hỗ trợ docx MIME type
                    # → Trích xuất toàn bộ text (bao gồm equations) và gửi dạng text
                    docx_text = self._extract_docx_text(file_path)
                    if docx_text:
                        parts.append(types.Part.from_text(
                            text=f"\n--- BẮT ĐẦU FILE: {label} ---\n{docx_text}\n--- KẾT THÚC FILE: {label} ---\n"
                        ))
                        print(f"  📄 Đã chuyển {label} thành text ({len(docx_text):,} chars)")
            except Exception as e:
                print(f"❌ Lỗi đọc file {file_path}: {e}")
                raise
        return parts

    @staticmethod
    def _extract_docx_text(filepath):
        """
        Trích xuất toàn bộ nội dung từ file docx, bao gồm equations (OMML).
        Dùng khi cần gửi docx cho AI mà Vertex AI không hỗ trợ docx MIME type.
        """
        try:
            # Thử import hàm extract từ check_de (có hỗ trợ equations)
            from check_de import extract_all_text_from_docx
            lines = extract_all_text_from_docx(filepath)
            return '\n'.join(lines)
        except ImportError:
            pass
        
        # Fallback: dùng python-docx cơ bản
        try:
            from docx import Document
            doc = Document(filepath)
            lines = []
            for para in doc.paragraphs:
                if para.text.strip():
                    lines.append(para.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            lines.append(cell.text.strip())
            return '\n'.join(lines)
        except Exception as e:
            print(f"  ⚠️  Không thể trích xuất text từ {filepath}: {e}")
            return None

    def send_data_to_AI(self, prompt, file_paths=None, temperature=0.12, top_p=0.8,
                        response_schema=None, max_output_tokens=65535,
                        response_mime_type=None, use_file_api=False):
        """
        Gửi request đến Gemini.

        Args:
            use_file_api: True → upload file qua File API (tái dùng trong session).
                          False (default) → nhúng inline như cũ.
        """
        if not self.client:
            return "❌ Lỗi: Client chưa được khởi tạo."

        user_parts = []

        if file_paths:
            if isinstance(file_paths, str):
                file_paths = [file_paths]

            if use_file_api:
                user_parts.extend(self.upload_files_cached(file_paths))
            else:
                user_parts.extend(self._inline_parts(file_paths))

        # Prompt text — log size để phát hiện request quá lớn
        prompt_chars = len(prompt)
        if prompt_chars > 800_000:
            print(f"  ⚠️  Prompt lớn: {prompt_chars:,} ký tự (~{prompt_chars//4:,} token ước tính)")
        user_parts.append(types.Part.from_text(text=prompt))

        contents = [types.Content(role="user", parts=user_parts)]

        config_args = {
            "temperature": temperature,
            "top_p": top_p,
            "max_output_tokens": max_output_tokens,
        }
        if response_schema:
            config_args["response_mime_type"] = "application/json"
            config_args["response_schema"] = response_schema
        elif response_mime_type:
            config_args["response_mime_type"] = response_mime_type

        generate_config = types.GenerateContentConfig(**config_args)

        try:
            response = self.client.models.generate_content(
                model=self.model_name,
                contents=contents,
                config=generate_config,
            )
            if response.text:
                return response.text
            else:
                return "⚠️ API trả về rỗng (Có thể do Safety Filter chặn)."
        except Exception as e:
            print(f"❌ Lỗi khi gọi AI generate_content: {e}")
            raise e